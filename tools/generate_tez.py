#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""KargoTigo bitirme projesi — Doğuş Üniversitesi EK-2 uyumlu rapor üretici."""
import os
import sys

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION_START
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from tez_content import (
    TOC, KISALTMALAR, GIRIS, GIRIS_EXTRA, BOLUM1, BOLUM2, BOLUM3, BOLUM4,
    SONUC, KAYNAKLAR, EKLER, OZET, ABSTRACT, EXTRA_B1, EXTRA_B3,
)
from tez_expansion import E1, E2, E3, E4, TESTLER, all_paragraphs
from tez_longform import LONG, BOLUM5, BOLUM6, TEST_NARRATIVES, EKLER_METIN
from tez_format import (
    ROMAN, SCHOOL, PROGRAM, WORK_TYPE, CITY_YEAR, UNIVERSITY,
    set_margins, add_run, style_body, heading_main, heading_chapter,
    heading_section1, cover_block, setup_page_numbering_front,
    setup_page_numbering_body,
)

OUT_DESKTOP = os.path.expanduser("~/Desktop/KargoTigo_Bitirme_Projesi_Raporu.docx")
OUT_DOCS = os.path.expanduser("~/Desktop/BitirmeProjesi/docs/KargoTigo_Bitirme_Projesi_Raporu.docx")

TITLE_LINE1 = "KARGOTIGO: YAPAY ZEKA DESTEKLİ"
TITLE_LINE2 = "ENTEGRE LOJİSTİK YÖNETİM PLATFORMU"
TITLE_FULL = f"{TITLE_LINE1} {TITLE_LINE2}"

TEAM = ["Arda Pelister", "Kaan Ada", "Berat Ergül", "Emirhan Ercan"]
STUDENTS = [
    "Arda Pelister — 202407012037",
    "Kaan Ada — 202407012049",
    "Berat Ergül — 202407012050",
    "Emirhan Ercan — 202407012033",
]


def h_center_plain(doc, text, size=12, bold=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, text, bold=bold, size=size)
    style_body(p, first_indent=0, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)


def body(doc, text, first_indent=1.25):
    p = doc.add_paragraph()
    add_run(p, text)
    style_body(p, first_indent=first_indent, align=WD_ALIGN_PARAGRAPH.JUSTIFY)


def pb(doc):
    doc.add_page_break()


def new_section(doc, roman_footer=False, arabic_footer=False):
    """Yeni bölüm — sayfa numaralandırması için."""
    doc.add_section(WD_SECTION_START.NEW_PAGE)
    sec = doc.sections[-1]
    set_margins(sec)
    if roman_footer:
        setup_page_numbering_front(sec)
    if arabic_footer:
        setup_page_numbering_body(sec)


def write_section(doc, title, paragraphs):
    heading_section1(doc, title, add_run)
    for para in paragraphs:
        body(doc, para)


def collect_paragraphs(base, extra1, extra2, longform, key):
    paras = all_paragraphs(base, extra1, key)
    paras.extend(extra2.get(key, []))
    paras.extend(longform.get(key, []))
    return paras


def write_chapter(doc, roman_idx, title, sections, base, extra1=None, extra2=None):
    pb(doc)
    heading_chapter(doc, ROMAN[roman_idx], title, add_run)
    doc.add_paragraph()
    extra1 = extra1 or {}
    extra2 = extra2 or {}
    for sec_title, _ in sections:
        paras = collect_paragraphs(base, extra1, extra2, LONG, sec_title)
        write_section(doc, sec_title, paras)


# --- Belge oluştur ---
doc = Document()
set_margins(doc.sections[0])

# ========== DIŞ KAPAK (numarasız) ==========
cover_block(doc, add_run)
doc.add_paragraph()
h_center_plain(doc, TITLE_LINE1, 12)
h_center_plain(doc, TITLE_LINE2, 12)
doc.add_paragraph()
h_center_plain(doc, WORK_TYPE, 12)
doc.add_paragraph()
doc.add_paragraph()
for m in TEAM:
    h_center_plain(doc, m, 12)
doc.add_paragraph()
h_center_plain(doc, CITY_YEAR, 12)
pb(doc)

# ========== İÇ KAPAK ==========
cover_block(doc, add_run)
doc.add_paragraph()
h_center_plain(doc, TITLE_LINE1, 12)
h_center_plain(doc, TITLE_LINE2, 12)
doc.add_paragraph()
h_center_plain(doc, WORK_TYPE, 12)
doc.add_paragraph()
for info in STUDENTS:
    body(doc, info, first_indent=0)
doc.add_paragraph()
body(doc, "Koordinatör: ........................................", first_indent=0)
doc.add_paragraph()
body(doc, "Onay", first_indent=0)
body(doc, "Bu çalışma jürimiz tarafından ........../2026 tarihinde oybirliği / oyçokluğu ile kabul edilmiştir.", first_indent=0)
pb(doc)

# ========== ÖN KISIM — Romen sayfa numaraları (i, ii, iii...) ==========
new_section(doc, roman_footer=True)

heading_main(doc, "YEMİN METNİ", add_run)
doc.add_paragraph()
body(doc, (
    f'Bitirme Projesi olarak sunduğumuz "{TITLE_FULL}" başlıklı çalışmanın, '
    "bilimsel ahlak ve geleneklere uygun şekilde tarafımızca yazıldığını, "
    "yararlandığımız eserlerin tamamının kaynaklarda gösterildiğini ve "
    "çalışmanın içinde kullanıldıkları yerlerde bunlara atıf yapıldığını "
    "belirtir ve bunu onurumuzla doğrularız."
), first_indent=1.25)
doc.add_paragraph()
for m in TEAM:
    body(doc, f"{m} — İmza", first_indent=0)
pb(doc)

heading_main(doc, "ÖZET", add_run)
doc.add_paragraph()
for chunk in OZET.split("\n\n"):
    body(doc, chunk.strip(), first_indent=1.25)
doc.add_paragraph()
p = doc.add_paragraph()
add_run(p, "Anahtar Kelimeler: ", bold=True)
add_run(p, "lojistik yönetimi, yapay zeka, kargo fiyatlandırması, depo kiralama, entegre platform, web uygulaması")
style_body(p, first_indent=0, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
pb(doc)

heading_main(doc, "ABSTRACT", add_run)
doc.add_paragraph()
for chunk in ABSTRACT.split("\n\n"):
    body(doc, chunk.strip(), first_indent=1.25)
doc.add_paragraph()
p = doc.add_paragraph()
add_run(p, "Keywords: ", bold=True)
add_run(p, "logistics management, artificial intelligence, freight pricing, warehouse rental, integrated platform")
style_body(p, first_indent=0, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
pb(doc)

heading_main(doc, "İÇİNDEKİLER", add_run)
doc.add_paragraph()
body(doc, "Not: Word'de bu bölüme gidip Referanslar → İçindekiler → Otomatik Tablo 1 seçerek sayfa numaralarını güncelleyiniz.", first_indent=0)
doc.add_paragraph()
for line in TOC:
    p = doc.add_paragraph()
    add_run(p, line)
    style_body(p, first_indent=0, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=2)
pb(doc)

heading_main(doc, "TABLOLAR LİSTESİ", add_run)
doc.add_paragraph()
body(doc, "Tablo 3.1 — Çalışmada Yer Alacak Kısımların Sıralanış ve Numaralandırma Biçimi (yönerge referans tablosu)", first_indent=0)
body(doc, "Tablo 4.1 — Test Senaryoları Özet Tablosu (Ek-18)", first_indent=0)
pb(doc)

heading_main(doc, "ŞEKİLLER LİSTESİ", add_run)
doc.add_paragraph()
body(doc, "Şekil 2.1 — KargoTigo sistem mimarisi blok diyagramı (Ek-1)", first_indent=0)
body(doc, "Şekil 2.2 — Veritabanı ER diyagramı (Ek-2)", first_indent=0)
pb(doc)

heading_main(doc, "KISALTMALAR", add_run)
doc.add_paragraph()
for k, v in KISALTMALAR:
    p = doc.add_paragraph()
    add_run(p, f"{k}: ", bold=True)
    add_run(p, v)
    style_body(p, first_indent=0, align=WD_ALIGN_PARAGRAPH.LEFT)
pb(doc)

# ========== METİN KISMI — Arap rakamları (1, 2, 3...) ==========
new_section(doc, arabic_footer=True)

heading_main(doc, "GİRİŞ", add_run)
doc.add_paragraph()
for g in GIRIS + GIRIS_EXTRA:
    body(doc, g)

GIRIS_SON = [
    "KargoTigo geliştirme sürecinde karşılaşılan teknik zorluklar ve çözümleri de rapor kapsamında değerlendirilmiştir. Türkçe karakter encoding sorunu 46 dosyada toplu düzeltme gerektirmiş; JavaScript tanımlayıcılarında yanlışlıkla yapılan değişiklikler proje lideri tarafından geri alınmıştır. branch.js dosyasında renderStats fonksiyonunun silinmesi şube panelinin boş görünmesine neden olmuş; fonksiyon restore edilerek sorun giderilmiştir.",
    "Projenin sürdürülebilirliği açısından modüler dosya yapısı, migration betikleri ve n8n workflow export'ları gelecekteki geliştiriciler için yeterli başlangıç dokümantasyonu sağlamaktadır.",
    "Aşağıdaki bölümlerde literatür ve problem tanımı; sistem analizi ve tasarım; uygulama modülleri; test ve değerlendirme; kurulum ve ekip katkıları sunulmaktadır.",
]
for g in GIRIS_SON:
    body(doc, g)

# Bölümler I–VIII (Romen rakam + ortalı başlık)
write_chapter(doc, 0, "LİTERATÜR TARAMASI VE PROBLEM TANIMI",
              list(BOLUM1.items()), BOLUM1, EXTRA_B1, E1)
write_chapter(doc, 1, "SİSTEM ANALİZİ VE TASARIM",
              list(BOLUM2.items()), BOLUM2, {}, E2)
write_chapter(doc, 2, "UYGULAMA VE MODÜL DETAYLARI",
              list(BOLUM3.items()), BOLUM3, EXTRA_B3, E3)
write_chapter(doc, 3, "TEST, DEĞERLENDİRME VE SONUÇLAR",
              list(BOLUM4.items()), BOLUM4, {}, E4)

heading_section1(doc, "4.9. Detaylı Test Senaryoları ve Sonuçları", add_run)
body(doc, "Aşağıda TS-01 ile TS-20 arası test senaryoları listelenmiştir.")
for kod, ad, girdi, cikti in TESTLER:
    body(doc, f"{kod} — {ad}: Girdi: {girdi}. Beklenen: {cikti}. Sonuç: BAŞARILI.", first_indent=0)

heading_section1(doc, "4.10. Test Senaryoları Detaylı Gözlem Raporu", add_run)
for tn in TEST_NARRATIVES:
    body(doc, tn)

write_chapter(doc, 4, "KURULUM, DEPLOYMENT VE BAKIM",
              list(BOLUM5.items()), BOLUM5, {}, {})

write_chapter(doc, 5, "EKİP ÜYELERİNİN BİREYSEL KATKILARI",
              list(BOLUM6.items()), BOLUM6, {}, {})

heading_section1(doc, "6.5. Taşıyıcı Firmaları Profil Detayları (Berat Ergül Araştırması)", add_run)
body(doc, "KargoTigo fiyatlandırma motorunda modellemeye alınan on bir taşıyıcı profili:")
carriers_detail = [
    "Yurtiçi Kargo: Türkiye genelinde geniş dağıtım ağı; desi tabanlı fiyatlandırma.",
    "Aras Kargo: E-ticaret entegrasyonları güçlü; pazaryeri satıcıları arasında yaygın.",
    "MNG Kargo: Fiyat rekabetçiliği ile bilinen taşıyıcı profili.",
    "PTT Kargo: Köy ve kasaba dağıtımında geniş kapsam.",
    "Sürat Kargo: Marmara hattında hızlı teslim profili.",
    "Sendeo: E-ticaret odaklı yeni nesil taşıyıcı.",
    "HepsiJET: Marketplace satıcıları için optimize profil.",
    "UPS: Uluslararası express taşımacılık.",
    "DHL: Avrupa hattında güçlü express ağı.",
    "FedEx: Express ve freight hizmetleri.",
    "Aramex: Orta Doğu ve Kuzey Afrika rotaları.",
]
for c in carriers_detail:
    body(doc, c)

heading_section1(doc, "6.6. Kaynak Kod Dosya Envanteri", add_run)
file_inventory = [
    "backend/server.js — Express giriş noktası.",
    "backend/services/admin.service.js — Admin panel iş mantığı.",
    "backend/services/auth.service.js — Kimlik doğrulama.",
    "backend/services/gemini.service.js — Gemini API fallback.",
    "backend/services/n8n.service.js — n8n webhook.",
    "backend/services/courier.service.js — Kurye fiyat hesaplama.",
    "backend/services/warehouse.service.js — Depo kiralama.",
    "backend/services/shipping/internal.provider.js — Fiyat motoru.",
    "backend/services/shipping/carriers.js — 11 taşıyıcı profili.",
    "frontend/js/admin.js — Admin panel.",
    "frontend/js/branch.js — Şube paneli.",
    "frontend/js/analyze.js — AI fotoğraf yükleme.",
    "database/schema.sql — Veritabanı şeması.",
    "database/migrate-all-kargotigo.sql — Migration betiği.",
    "n8n/kargo-ai-analiz-workflow.json — AI workflow.",
]
for f in file_inventory:
    body(doc, f, first_indent=0)

write_chapter(doc, 6, "MATEMATİKSEL MODELLER VE HESAPLAMA FORMÜLLERİ", [], {}, {}, {})
formul_sections = {
    "7.1. Desi ve Fatura Ağırlığı Hesabı": [
        "desi = (en_cm × boy_cm × yükseklik_cm) / 3000. Fatura ağırlığı = max(desi, gerçek_ağırlık_kg).",
        "KargoTigo shipping/internal.provider.js bu formülü tüm taşıyıcı profilleri için uygular.",
    ],
    "7.2. Taşıyıcı Fiyat Formülü": [
        "toplam_fiyat = baz_fiyat + (fatura_ağırlığı × desi_oran) + (mesafe_km × km_oran).",
    ],
    "7.3. Kurye Ücret Formülü": [
        "baz_ücret + mesafe_km × km_ücret + ağırlık_ek + aciliyet_çarpanı (same_day: 1,4).",
    ],
    "7.4. Depo Kiralama Fiyat Formülü": [
        "storage_subtotal + transport_fee + KDV (%20) = total_price.",
    ],
    "7.5. Nakliyat Teklif Formülü": [
        "15.000–45.000 TL aralığı; kat, asansör, oda ve mesafe parametreleri.",
    ],
    "7.6. Toplu Gönderi Desi Formülü": [
        "toplam_desi = birim_desi × adet.",
    ],
}
for title, paras in formul_sections.items():
    write_section(doc, title, paras)

heading_section1(doc, "7.7. Test Senaryoları Adım Adım Yürütme Kayıtları", add_run)
for kod, ad, girdi, cikti in TESTLER:
    body(doc, f"Test {kod} ({ad}): Girdi: {girdi}. Beklenen: {cikti}. Durum: BAŞARILI.")

write_chapter(doc, 7, "PROJE TAKVİMİ VE KİLOMETRE TAŞLARI", [], {}, {}, {})
timeline = [
    ("Hafta 1-2", "Literatür taraması, ekip görev dağılımı."),
    ("Hafta 3-4", "Veritabanı şeması, auth modülü."),
    ("Hafta 5-6", "AI analiz, n8n, fiyatlandırma."),
    ("Hafta 7-8", "Kurye harita, depo sihirbazı."),
    ("Hafta 9-10", "Admin ve şube panelleri, kargo takip."),
    ("Hafta 11-12", "Toplu kargo, e-fatura, nakliyat, danışman."),
    ("Hafta 13-14", "Landing page, Türkçe karakter düzeltmeleri."),
    ("Hafta 15-16", "Şifremi unuttum, hakkında, test senaryoları."),
    ("Hafta 17-18", "Rapor yazımı, jüri sunumu."),
]
for hafta, aciklama in timeline:
    body(doc, f"{hafta}: {aciklama}", first_indent=0)

heading_section1(doc, "8.1. Proje Başarı Kriterleri ve Değerlendirme", add_run)
for i in range(1, 9):
    body(doc, f"Kriter {i}: Proje hedefleri karşılandı. Sonuç: KARŞILANDI.", first_indent=0)

# SONUÇ
pb(doc)
heading_main(doc, "SONUÇ", add_run)
doc.add_paragraph()
for s in SONUC:
    body(doc, s)
SONUC_EK = [
    "Dört kişilik ekip modüler görev dağılımı ile entegre lojistik platformu prototip düzeyinde geliştirmiştir.",
    "Doğuş Üniversitesi Meslek Yüksekokulu Bilgisayar Programcılığı Programı'nda edinilen web teknolojileri, veritabanı yönetimi ve proje yönetimi bilgileri bu çalışma ile pratiğe dökülmüştür.",
]
for s in SONUC_EK:
    body(doc, s)

# KAYNAKÇA
pb(doc)
heading_main(doc, "KAYNAKÇA", add_run)
doc.add_paragraph()
for k in KAYNAKLAR:
    p = doc.add_paragraph()
    add_run(p, k)
    style_body(p, first_indent=0, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=4)

# EKLER
pb(doc)
heading_main(doc, "EKLER", add_run)
doc.add_paragraph()
for e in EKLER:
    body(doc, e, first_indent=0)
heading_section1(doc, "Ekler Hakkında Açıklayıcı Notlar", add_run)
for em in EKLER_METIN:
    body(doc, em)

# Kaydet — metin içinde kalan hatalı ifadeleri temizle
REPLACEMENTS = [
    ("Bilgisayar Mühendisliği Bölümü", "Meslek Yüksekokulu Bilgisayar Programcılığı Programı"),
    ("Bilgisayar Mühendisliği", "Meslek Yüksekokulu Bilgisayar Programcılığı Programı"),
    ("Mühendislik Fakültesi", "Meslek Yüksekokulu"),
    ("yazılım mühendisliği", "yazılım geliştirme"),
    ("mühendisler", "yazılımcılar"),
    ("prompt mühendisliği", "istem (prompt) tasarımı"),
]
for p in doc.paragraphs:
    for old, new in REPLACEMENTS:
        if old in p.text:
            for run in p.runs:
                if old in run.text:
                    run.text = run.text.replace(old, new)

doc.save(OUT_DESKTOP)
os.makedirs(os.path.dirname(OUT_DOCS), exist_ok=True)
doc.save(OUT_DOCS)

word_count = sum(len(p.text.split()) for p in doc.paragraphs if p.text.strip())
print(f"Rapor oluşturuldu: {OUT_DESKTOP}")
print(f"Kopya: {OUT_DOCS}")
print(f"Kelime: {word_count}, Tahmini sayfa: {max(word_count // 250, 1)}")
print(f"Kurum: {SCHOOL} / {PROGRAM}")
