# -*- coding: utf-8 -*-
"""Doğuş Üniversitesi EK-2 yazım kılavuzu biçim yardımcıları."""
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Kurumsal bilgiler
UNIVERSITY = "DOĞUŞ ÜNİVERSİTESİ"
SCHOOL = "MESLEK YÜKSEKOKULU"
PROGRAM = "BİLGİSAYAR PROGRAMCILIĞI PROGRAMI"
WORK_TYPE = "Bitirme Projesi"
CITY_YEAR = "İSTANBUL, 2026"

ROMAN = ["I", "II", "III", "IV", "V", "VI", "VII", "VIII", "IX", "X"]


def set_margins(section):
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(4)
    section.right_margin = Cm(2.5)


def _font(run, size=12, bold=False):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    for attr in ("ascii", "hAnsi", "cs"):
        rFonts.set(qn(f"w:{attr}"), "Times New Roman")
    rPr.insert(0, rFonts)


def add_run(p, text, bold=False, size=12):
    run = p.add_run(text)
    _font(run, size, bold)
    return run


def style_body(p, first_indent=1.25, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=6):
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    p.paragraph_format.space_after = Pt(space_after)
    if first_indent:
        p.paragraph_format.first_line_indent = Cm(first_indent)
    p.alignment = align
    for run in p.runs:
        _font(run)


def heading_main(doc, text, add_run_fn):
    """GİRİŞ, ÖZET, SONUÇ vb. — ortalı, büyük harf."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run_fn(p, text.upper(), bold=True, size=12)
    style_body(p, first_indent=0, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)


def heading_chapter(doc, roman, title, add_run_fn):
    """Bölüm: Romen rakam + ortalı başlık (iki satır)."""
    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run_fn(p1, roman, bold=True, size=12)
    style_body(p1, first_indent=0, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run_fn(p2, title.upper(), bold=True, size=12)
    style_body(p2, first_indent=0, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)


def heading_section1(doc, title, add_run_fn):
    """1.1. — sola dayalı, birinci derece, tamamı büyük harf."""
    p = doc.add_paragraph()
    add_run_fn(p, title.upper(), bold=True, size=12)
    style_body(p, first_indent=0, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=6)


def heading_section2(doc, title, add_run_fn):
    """1.1.1 — sola dayalı, kelime başları büyük."""
    p = doc.add_paragraph()
    add_run_fn(p, title, bold=True, size=12)
    style_body(p, first_indent=0, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=6)


def cover_block(doc, add_run_fn, inner=False):
    """Dış / iç kapak ortak üst blok."""
    for line, size in [
        ("T. C.", 12),
        (UNIVERSITY, 12),
        (SCHOOL, 12),
        (PROGRAM, 12),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run_fn(p, line, bold=False, size=size)
        style_body(p, first_indent=0, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)


def _field_run(parent, field_code):
    """Word alan kodu (PAGE vb.)."""
    run = parent.add_run()
    r = run._r
    for part, text in [("begin", None), ("instrText", field_code), ("separate", None), ("end", None)]:
        el = OxmlElement("w:fldChar")
        el.set(qn("w:fldCharType"), part if part != "instrText" else "begin")
        if part == "instrText":
            el = OxmlElement("w:instrText")
            el.set(qn("xml:space"), "preserve")
            el.text = field_code
        elif part == "separate":
            el = OxmlElement("w:fldChar")
            el.set(qn("w:fldCharType"), "separate")
        elif part == "end":
            el = OxmlElement("w:fldChar")
            el.set(qn("w:fldCharType"), "end")
        r.append(el)
    return run


def _footer_page_number(section, roman=False):
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Temizle
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    fmt = " \\* ROMAN \\* MERGEFORMAT" if roman else " \\* MERGEFORMAT"
    run = p.add_run()
    r = run._r
    fc1 = OxmlElement("w:fldChar"); fc1.set(qn("w:fldCharType"), "begin"); r.append(fc1)
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve")
    instr.text = f" PAGE {fmt.strip()} "
    r.append(instr)
    fc2 = OxmlElement("w:fldChar"); fc2.set(qn("w:fldCharType"), "separate"); r.append(fc2)
    fc3 = OxmlElement("w:fldChar"); fc3.set(qn("w:fldCharType"), "end"); r.append(fc3)
    _font(run, 12, False)


def add_section_break(doc):
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    sectPr = OxmlElement("w:sectPr")
    type_el = OxmlElement("w:type")
    type_el.set(qn("w:val"), "nextPage")
    sectPr.append(type_el)
    pPr.append(sectPr)


def setup_page_numbering_front(section):
    """Ön kısım: küçük Romen rakamları (i, ii, iii...)."""
    sectPr = section._sectPr
    pgNumType = OxmlElement("w:pgNumType")
    pgNumType.set(qn("w:fmt"), "lowerRoman")
    pgNumType.set(qn("w:start"), "1")
    sectPr.append(pgNumType)
    _footer_page_number(section, roman=True)


def setup_page_numbering_body(section):
    """Metin kısmı: Arap rakamları 1, 2, 3..."""
    sectPr = section._sectPr
    pgNumType = OxmlElement("w:pgNumType")
    pgNumType.set(qn("w:fmt"), "decimal")
    pgNumType.set(qn("w:start"), "1")
    sectPr.append(pgNumType)
    _footer_page_number(section, roman=False)
